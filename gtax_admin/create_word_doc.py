"""
Script to convert DOCUMENTATION.md to a Word document.
Requires: pip install python-docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import re


def create_trading_platform_doc():
    """Create a Word document from the documentation."""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 102, 153)
    
    # =====================================================
    # TITLE PAGE
    # =====================================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title_run = title.add_run("AI-Powered Autonomous Trading Platform")
    title_run.bold = True
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Complete Technical Documentation")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    version = doc.add_paragraph()
    version_run = version.add_run("Version 1.0")
    version_run.font.size = Pt(14)
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph()
    date_run = date.add_run("March 2026")
    date_run.font.size = Pt(14)
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # =====================================================
    # TABLE OF CONTENTS
    # =====================================================
    doc.add_heading("Table of Contents", level=1)
    
    toc_items = [
        ("1. Overview", 1),
        ("2. Architecture", 1),
        ("3. Technology Stack", 1),
        ("4. Project Structure", 1),
        ("5. Installation", 1),
        ("6. Configuration", 1),
        ("7. Workflow", 1),
        ("8. Modules Documentation", 1),
        ("9. API Reference", 1),
        ("10. Trade Logging", 1),
        ("11. Risk Management", 1),
        ("12. Troubleshooting", 1),
    ]
    
    for item, level in toc_items:
        p = doc.add_paragraph()
        p.add_run(item).font.size = Pt(12)
        p.paragraph_format.left_indent = Inches(0.25 * (level - 1))
    
    doc.add_page_break()
    
    # =====================================================
    # 1. OVERVIEW
    # =====================================================
    doc.add_heading("1. Overview", level=1)
    
    doc.add_paragraph(
        "This platform is designed to operate fully autonomously - from stock selection "
        "to trade execution. Once started, it continuously:"
    )
    
    overview_steps = [
        "Scans markets for trading opportunities",
        "Analyzes using AI (technical indicators + chart patterns + news sentiment)",
        "Generates signals with confidence scores",
        "Executes trades when confidence exceeds threshold",
        "Manages positions (stop-loss, take-profit)",
        "Logs everything to Excel for review"
    ]
    
    for i, step in enumerate(overview_steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)
    
    doc.add_heading("Key Features", level=2)
    
    features = [
        ("Autonomous Operation", "Runs 24/7 without human intervention"),
        ("AI Decision Engine", "Weighted scoring system combining multiple analysis methods"),
        ("Multi-Asset Support", "Stocks and Cryptocurrencies"),
        ("Risk Management", "Position sizing, stop-loss, take-profit"),
        ("Paper Trading", "Test strategies without real money"),
        ("Real Trading Ready", "Alpaca/Binance API integration"),
        ("Excel Trade Logging", "Complete trade history with analysis data"),
        ("Web Dashboard", "Real-time monitoring via React frontend"),
        ("WebSocket Updates", "Live price and signal updates"),
    ]
    
    for feature, description in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{feature}: ").bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # =====================================================
    # 2. ARCHITECTURE
    # =====================================================
    doc.add_heading("2. Architecture", level=1)
    
    doc.add_paragraph(
        "The system follows a layered architecture with the following components:"
    )
    
    doc.add_heading("Frontend Layer (React)", level=2)
    p = doc.add_paragraph()
    p.add_run("Components: ").bold = True
    p.add_run("Dashboard, Signals, Portfolio, Backtest, Stock Detail pages")
    
    doc.add_heading("Backend Layer (FastAPI)", level=2)
    p = doc.add_paragraph()
    p.add_run("Core Component: ").bold = True
    p.add_run("Autonomous Trading Orchestrator")
    
    doc.add_paragraph("The orchestrator coordinates the following flow:")
    
    flow_steps = [
        "Scan Markets → Analyze AI → Signal Generate → Execute Trade"
    ]
    for step in flow_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(step)
    
    doc.add_heading("Service Modules", level=2)
    
    modules = [
        ("Market Scanner", "Scans market for opportunities"),
        ("News Sentiment Engine", "Analyzes news sentiment using VADER/TextBlob"),
        ("Chart Pattern Detection", "Detects technical patterns"),
        ("AI Decision Engine", "Combines all analysis into signals"),
        ("Trade Executor", "Executes trades with risk management"),
        ("Learning Model", "Improves predictions over time"),
        ("Excel Logger", "Logs all trades to Excel"),
    ]
    
    for module, desc in modules:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{module}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("Data Layer", level=2)
    doc.add_paragraph("Database: SQLite (development) / PostgreSQL (production)")
    doc.add_paragraph("Cache: Redis for market data caching")
    doc.add_paragraph("Broker APIs: Alpaca (stocks), Binance via CCXT (crypto)")
    
    doc.add_page_break()
    
    # =====================================================
    # 3. TECHNOLOGY STACK
    # =====================================================
    doc.add_heading("3. Technology Stack", level=1)
    
    doc.add_heading("Backend Libraries", level=2)
    
    # Create table for backend libraries
    backend_libs = [
        ("Component", "Library", "Version", "Purpose"),
        ("Web Framework", "FastAPI", "≥0.104.0", "Async REST API with automatic OpenAPI docs"),
        ("ASGI Server", "Uvicorn", "≥0.24.0", "High-performance async server"),
        ("Database ORM", "SQLAlchemy", "≥2.0.0", "Async database operations"),
        ("Migrations", "Alembic", "≥1.12.0", "Database schema migrations"),
        ("PostgreSQL Driver", "asyncpg", "≥0.29.0", "Async PostgreSQL support"),
        ("SQLite Driver", "aiosqlite", "≥0.19.0", "Async SQLite for development"),
        ("Caching", "Redis", "≥5.0.0", "Market data caching"),
        ("Background Tasks", "Celery", "≥5.3.0", "Distributed task queue"),
        ("Logging", "Loguru", "≥0.7.0", "Structured logging"),
        ("Config", "Pydantic-Settings", "≥2.1.0", "Environment configuration"),
    ]
    
    table = doc.add_table(rows=len(backend_libs), cols=4)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(backend_libs):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data
            if i == 0:  # Header row
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Data & Analysis Libraries", level=2)
    
    data_libs = [
        ("Component", "Library", "Version", "Purpose"),
        ("Data Processing", "Pandas", "≥2.1.0", "DataFrame operations"),
        ("Numerical", "NumPy", "≥1.26.0", "Array computations"),
        ("Scientific", "SciPy", "≥1.11.0", "Statistical functions, peak detection"),
        ("Technical Analysis", "TA", "≥0.11.0", "Technical indicators (RSI, MACD, etc.)"),
        ("TA-Lib", "ta-lib", "≥0.4.28", "Advanced technical analysis"),
        ("Market Data", "yfinance", "≥0.2.31", "Yahoo Finance API for stock/crypto data"),
        ("Crypto Data", "CCXT", "≥4.1.0", "Cryptocurrency exchange connections"),
    ]
    
    table = doc.add_table(rows=len(data_libs), cols=4)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data_libs):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("AI & Machine Learning Libraries", level=2)
    
    ai_libs = [
        ("Component", "Library", "Version", "Purpose"),
        ("ML Framework", "scikit-learn", "≥1.3.0", "Classification, regression algorithms"),
        ("Deep Learning", "TensorFlow", "≥2.15.0", "Neural networks (optional)"),
        ("Gradient Boosting", "XGBoost", "≥2.0.0", "Gradient boosting trees"),
        ("Light GBM", "LightGBM", "≥4.1.0", "Fast gradient boosting"),
        ("NLP", "NLTK", "≥3.8.0", "Natural language processing toolkit"),
        ("NLP Models", "spaCy", "≥3.7.0", "Industrial NLP"),
        ("Transformers", "transformers", "≥4.35.0", "Pre-trained language models"),
        ("Sentiment", "VADER", "≥3.3.2", "Social media/financial sentiment"),
        ("Sentiment", "TextBlob", "≥0.17.0", "Simple sentiment analysis"),
    ]
    
    table = doc.add_table(rows=len(ai_libs), cols=4)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(ai_libs):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Trading & Execution Libraries", level=2)
    
    trading_libs = [
        ("Component", "Library", "Purpose"),
        ("Stock Broker", "Alpaca API", "Paper/live stock trading"),
        ("Crypto Exchange", "Binance API (via CCXT)", "Cryptocurrency trading"),
        ("HTTP Client", "httpx (≥0.25.0)", "Async HTTP requests"),
        ("News Feeds", "feedparser (≥6.0.0)", "RSS news parsing"),
    ]
    
    table = doc.add_table(rows=len(trading_libs), cols=3)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(trading_libs):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Frontend Libraries", level=2)
    
    frontend_libs = [
        ("Component", "Library", "Version", "Purpose"),
        ("Framework", "React", "18.x", "UI component framework"),
        ("Language", "TypeScript", "5.x", "Type safety"),
        ("Styling", "TailwindCSS", "3.x", "Utility-first CSS framework"),
        ("Charts", "Chart.js", "4.x", "Data visualization"),
        ("Charts", "react-chartjs-2", "5.x", "React Chart.js wrapper"),
        ("HTTP", "Axios", "1.x", "API requests"),
        ("Routing", "React Router", "6.x", "Client-side routing"),
        ("Icons", "Lucide React", "-", "Icon library"),
        ("Date", "date-fns", "2.x", "Date formatting utilities"),
    ]
    
    table = doc.add_table(rows=len(frontend_libs), cols=4)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(frontend_libs):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Logging & Export", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("openpyxl (≥3.1.0): ").bold = True
    p.add_run("Trade log Excel files with formatting")
    
    doc.add_page_break()
    
    # =====================================================
    # 4. PROJECT STRUCTURE
    # =====================================================
    doc.add_heading("4. Project Structure", level=1)
    
    doc.add_paragraph("The project is organized as follows:")
    
    structure = """gtax_admin/
├── backend/
│   ├── app/
│   │   ├── main.py                    # FastAPI application entry point
│   │   ├── api/                       # API Endpoints
│   │   │   ├── auto_trading.py        # Autonomous trading control
│   │   │   ├── backtest.py            # Backtesting endpoints
│   │   │   ├── patterns.py            # Chart pattern detection
│   │   │   ├── scanner.py             # Market scanner endpoints
│   │   │   ├── sentiment.py           # News sentiment endpoints
│   │   │   ├── signals.py             # Trading signals endpoints
│   │   │   ├── trade_log.py           # Excel log endpoints
│   │   │   ├── trades.py              # Trade execution endpoints
│   │   │   └── websocket.py           # Real-time WebSocket
│   │   ├── core/
│   │   │   └── config.py              # Application settings
│   │   ├── db/
│   │   │   ├── database.py            # Database connection
│   │   │   └── models.py              # SQLAlchemy ORM models
│   │   └── services/
│   │       ├── ai/decision_engine.py       # AI Decision Engine
│   │       ├── data/data_fetcher.py        # Market data service
│   │       ├── learning/model_trainer.py   # ML model trainer
│   │       ├── logging/trade_logger.py     # Excel trade logger
│   │       ├── orchestrator/trading_orchestrator.py  # Main loop
│   │       ├── patterns/pattern_detector.py # Pattern detection
│   │       ├── scanner/market_scanner.py   # Opportunity scanner
│   │       ├── sentiment/news_analyzer.py  # Sentiment analysis
│   │       └── trading/trade_executor.py   # Order execution
│   ├── logs/trades/                   # Excel trade logs
│   └── requirements.txt               # Python dependencies
├── frontend/
│   ├── src/
│   │   ├── components/                # Reusable UI components
│   │   ├── pages/                     # Page components
│   │   ├── services/                  # API services
│   │   └── types/                     # TypeScript definitions
│   └── package.json
├── docker-compose.yml
├── DOCUMENTATION.md
└── README.md"""
    
    p = doc.add_paragraph()
    p.add_run(structure).font.name = 'Courier New'
    p.add_run().font.size = Pt(9)
    
    doc.add_page_break()
    
    # =====================================================
    # 5. INSTALLATION
    # =====================================================
    doc.add_heading("5. Installation", level=1)
    
    doc.add_heading("Prerequisites", level=2)
    prereqs = [
        "Python 3.10+",
        "Node.js 18+",
        "Redis (optional, for caching)",
        "PostgreSQL (optional, SQLite works for development)"
    ]
    for prereq in prereqs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(prereq)
    
    doc.add_heading("Backend Setup", level=2)
    
    backend_commands = """# Navigate to backend directory
cd backend

# Create virtual environment
python -m venv venv

# Activate virtual environment
# Windows:
venv\\Scripts\\activate
# Linux/Mac:
source venv/bin/activate

# Install dependencies
pip install -r requirements.txt

# Copy environment file
cp .env.example .env"""
    
    p = doc.add_paragraph()
    p.add_run(backend_commands).font.name = 'Courier New'
    
    doc.add_heading("Frontend Setup", level=2)
    
    frontend_commands = """# Navigate to frontend directory
cd frontend

# Install dependencies
npm install

# Start development server
npm start"""
    
    p = doc.add_paragraph()
    p.add_run(frontend_commands).font.name = 'Courier New'
    
    doc.add_heading("Running the Application", level=2)
    
    doc.add_paragraph("Option 1: Manual Start")
    
    run_commands = """# Terminal 1 - Backend
cd backend
uvicorn app.main:app --reload --host 0.0.0.0 --port 8000

# Terminal 2 - Frontend
cd frontend
npm start"""
    
    p = doc.add_paragraph()
    p.add_run(run_commands).font.name = 'Courier New'
    
    doc.add_paragraph()
    doc.add_paragraph("Option 2: Docker Compose")
    
    p = doc.add_paragraph()
    p.add_run("docker-compose up -d").font.name = 'Courier New'
    
    doc.add_paragraph()
    doc.add_paragraph("The application will be available at:")
    urls = [
        ("Frontend", "http://localhost:3000"),
        ("Backend API", "http://localhost:8000"),
        ("API Documentation", "http://localhost:8000/docs"),
    ]
    for name, url in urls:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{name}: ").bold = True
        p.add_run(url)
    
    doc.add_page_break()
    
    # =====================================================
    # 6. CONFIGURATION
    # =====================================================
    doc.add_heading("6. Configuration", level=1)
    
    doc.add_heading("Environment Variables (.env)", level=2)
    
    config_sections = [
        ("Application Settings", [
            ("APP_NAME", "AI Trading Platform"),
            ("DEBUG", "true"),
            ("SECRET_KEY", "your-secret-key"),
        ]),
        ("Database Configuration", [
            ("DATABASE_URL", "sqlite+aiosqlite:///./trading.db"),
        ]),
        ("Redis Cache", [
            ("REDIS_URL", "redis://localhost:6379/0"),
        ]),
        ("External API Keys", [
            ("NEWS_API_KEY", "your-newsapi-key"),
            ("ALPHA_VANTAGE_KEY", "your-alpha-vantage-key"),
        ]),
        ("Trading Broker APIs", [
            ("ALPACA_API_KEY", "your-alpaca-key"),
            ("ALPACA_SECRET_KEY", "your-alpaca-secret"),
            ("ALPACA_BASE_URL", "https://paper-api.alpaca.markets"),
            ("BINANCE_API_KEY", "your-binance-key"),
            ("BINANCE_SECRET_KEY", "your-binance-secret"),
        ]),
        ("AI Decision Engine", [
            ("AI_TECHNICAL_WEIGHT", "0.4 (40% weight to technical analysis)"),
            ("AI_PATTERN_WEIGHT", "0.3 (30% weight to chart patterns)"),
            ("AI_SENTIMENT_WEIGHT", "0.3 (30% weight to news sentiment)"),
            ("AI_CONFIDENCE_THRESHOLD", "0.65 (Minimum confidence to generate signal)"),
        ]),
        ("Autonomous Trading", [
            ("AUTO_START_TRADING", "true (Start trading on startup)"),
            ("AUTO_TRADING_INTERVAL", "300 (Scan interval in seconds)"),
            ("AUTO_MAX_POSITIONS", "5 (Maximum concurrent positions)"),
            ("AUTO_CONFIDENCE_THRESHOLD", "0.7 (Minimum confidence to execute)"),
            ("AUTO_TRADE_24_7", "true (Trade outside market hours)"),
        ]),
        ("Risk Management", [
            ("MAX_POSITION_SIZE", "0.1 (10% of portfolio per position)"),
            ("DEFAULT_STOP_LOSS", "0.02 (2% stop loss)"),
            ("DEFAULT_TAKE_PROFIT", "0.05 (5% take profit)"),
        ]),
    ]
    
    for section_name, settings in config_sections:
        doc.add_heading(section_name, level=3)
        table = doc.add_table(rows=len(settings) + 1, cols=2)
        table.style = 'Table Grid'
        
        # Header
        table.rows[0].cells[0].text = "Variable"
        table.rows[0].cells[1].text = "Default Value"
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
        
        for i, (var, val) in enumerate(settings, 1):
            table.rows[i].cells[0].text = var
            table.rows[i].cells[1].text = val
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # =====================================================
    # 7. WORKFLOW
    # =====================================================
    doc.add_heading("7. Workflow", level=1)
    
    doc.add_paragraph(
        "The system runs through these steps continuously in an autonomous loop:"
    )
    
    workflow_steps = [
        ("Step 1: Application Startup", 
         "FastAPI server starts, database initializes, AutomatedTradingOrchestrator begins.",
         "Libraries: FastAPI, SQLAlchemy, asyncio, Loguru"),
        
        ("Step 2: Market Scanning (Every 5 minutes)",
         "Fetches latest price data, calculates momentum/volume/volatility, ranks symbols.",
         "Libraries: yfinance, pandas, numpy, ta"),
        
        ("Step 3: Technical Analysis",
         "Calculates RSI, MACD, Bollinger Bands, Moving Averages, Volume, ATR for each symbol.",
         "Libraries: ta, ta-lib, pandas, numpy"),
        
        ("Step 4: Chart Pattern Detection",
         "Detects Double Bottom/Top, Head & Shoulders, Triangles, Flags, Support/Resistance.",
         "Libraries: numpy, scipy (argrelextrema)"),
        
        ("Step 5: News Sentiment Analysis",
         "Fetches news, runs VADER and TextBlob sentiment analysis, aggregates scores.",
         "Libraries: nltk, vaderSentiment, textblob, feedparser, httpx"),
        
        ("Step 6: AI Decision Engine",
         "Combines scores: (Technical × 0.40) + (Pattern × 0.30) + (Sentiment × 0.30). "
         "Generates BUY/SELL/HOLD signals with confidence scores.",
         "Libraries: numpy, scikit-learn"),
        
        ("Step 7: Trade Execution",
         "Validates trade (position limits, cash available), calculates position size, executes order.",
         "Libraries: httpx, ccxt, asyncio"),
        
        ("Step 8: Excel Trade Logging",
         "Logs trade details: ID, symbol, prices, strategy, confidence, P&L to Excel.",
         "Libraries: openpyxl"),
        
        ("Step 9: Position Monitoring (Every 30s)",
         "Checks stop-loss and take-profit triggers, auto-closes positions when triggered.",
         "Libraries: yfinance, asyncio"),
        
        ("Step 10: Cycle Repeats",
         "System waits for next scan interval and repeats the cycle 24/7.",
         "File: trading_orchestrator.py"),
    ]
    
    for step_name, description, libs in workflow_steps:
        doc.add_heading(step_name, level=2)
        doc.add_paragraph(description)
        p = doc.add_paragraph()
        p.add_run(libs).italic = True
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # =====================================================
    # 8. MODULES DOCUMENTATION
    # =====================================================
    doc.add_heading("8. Modules Documentation", level=1)
    
    modules_doc = [
        ("Market Scanner", 
         "services/scanner/market_scanner.py",
         "Continuously scans the market for trading opportunities.",
         ["scan_markets() - Scan all watchlist symbols",
          "get_top_opportunities(limit) - Get top ranked opportunities",
          "analyze_symbol(symbol) - Analyze specific symbol"]),
        
        ("AI Decision Engine",
         "services/ai/decision_engine.py",
         "The brain - combines all analysis into actionable signals.",
         ["analyze(symbol) - Full analysis with signal generation",
          "generate_signal(symbol) - Get trading signal only",
          "get_technical_score(data) - Technical analysis only"]),
        
        ("Chart Pattern Detector",
         "services/patterns/pattern_detector.py",
         "Identifies chart patterns that predict price movements.",
         ["detect_patterns(symbol) - Detect all patterns",
          "get_pattern_score(symbol) - Get pattern-based score"]),
        
        ("News Sentiment Analyzer",
         "services/sentiment/news_analyzer.py",
         "Analyzes news sentiment for trading decisions.",
         ["analyze_sentiment(symbol) - Get sentiment analysis",
          "fetch_news(symbol) - Get raw news articles"]),
        
        ("Trade Executor",
         "services/trading/trade_executor.py",
         "Executes trades with risk management.",
         ["execute_trade(symbol, side, risk_pct, ...) - Execute a trade",
          "get_portfolio() - Get portfolio summary",
          "update_and_check_orders() - Check stop loss / take profit"]),
        
        ("Trading Orchestrator",
         "services/orchestrator/trading_orchestrator.py",
         "Coordinates autonomous trading loop.",
         ["start() - Start autonomous trading",
          "stop() - Stop autonomous trading",
          "run_trading_cycle() - Run one trading cycle"]),
        
        ("Excel Trade Logger",
         "services/logging/trade_logger.py",
         "Maintains detailed Excel trade logs.",
         ["log_entry(...) - Log a trade entry",
          "log_exit(...) - Log a trade exit",
          "get_trade_summary() - Get summary statistics"]),
    ]
    
    for name, filepath, purpose, functions in modules_doc:
        doc.add_heading(name, level=2)
        p = doc.add_paragraph()
        p.add_run("File: ").bold = True
        p.add_run(filepath).font.name = 'Courier New'
        
        p = doc.add_paragraph()
        p.add_run("Purpose: ").bold = True
        p.add_run(purpose)
        
        p = doc.add_paragraph()
        p.add_run("Key Functions:").bold = True
        
        for func in functions:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(func).font.name = 'Courier New'
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # =====================================================
    # 9. API REFERENCE
    # =====================================================
    doc.add_heading("9. API Reference", level=1)
    
    api_sections = [
        ("Autonomous Trading Endpoints", [
            ("GET", "/api/auto/status", "Get orchestrator status"),
            ("POST", "/api/auto/start", "Start autonomous trading"),
            ("POST", "/api/auto/stop", "Stop autonomous trading"),
            ("POST", "/api/auto/configure", "Update configuration"),
            ("POST", "/api/auto/run-cycle", "Manually trigger one cycle"),
        ]),
        ("Trade Log Endpoints", [
            ("GET", "/api/log/summary", "Trade statistics (win rate, P&L)"),
            ("GET", "/api/log/open-trades", "Currently open positions"),
            ("GET", "/api/log/download", "Download Excel log file"),
            ("GET", "/api/log/info", "Log file information"),
        ]),
        ("Market Analysis Endpoints", [
            ("POST", "/api/scanner/scan", "Scan markets for opportunities"),
            ("GET", "/api/scanner/analyze/{symbol}", "Analyze specific symbol"),
            ("GET", "/api/sentiment/{symbol}", "Get sentiment analysis"),
            ("GET", "/api/patterns/{symbol}", "Detect chart patterns"),
            ("POST", "/api/signals/generate", "Generate trading signals"),
        ]),
        ("Trading Endpoints", [
            ("POST", "/api/trades/execute", "Execute a trade manually"),
            ("GET", "/api/trades/portfolio", "Get portfolio summary"),
            ("GET", "/api/trades/positions", "Get open positions"),
            ("GET", "/api/trades/history", "Get trade history"),
        ]),
        ("Backtesting Endpoints", [
            ("POST", "/api/backtest/run", "Run historical backtest"),
            ("GET", "/api/backtest/results/{id}", "Get backtest results"),
        ]),
        ("WebSocket Endpoints", [
            ("WS", "/ws/prices", "Real-time price updates"),
            ("WS", "/ws/signals", "Real-time signal notifications"),
            ("WS", "/ws/trades", "Trade execution notifications"),
        ]),
    ]
    
    for section_name, endpoints in api_sections:
        doc.add_heading(section_name, level=2)
        
        table = doc.add_table(rows=len(endpoints) + 1, cols=3)
        table.style = 'Table Grid'
        
        # Header
        headers = ["Method", "Endpoint", "Description"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for i, (method, endpoint, desc) in enumerate(endpoints, 1):
            table.rows[i].cells[0].text = method
            table.rows[i].cells[1].text = endpoint
            table.rows[i].cells[2].text = desc
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # =====================================================
    # 10. TRADE LOGGING
    # =====================================================
    doc.add_heading("10. Trade Logging", level=1)
    
    doc.add_heading("Excel File Structure", level=2)
    p = doc.add_paragraph()
    p.add_run("File Location: ").bold = True
    p.add_run("logs/trades/trade_log_YYYY_MM.xlsx")
    
    doc.add_paragraph("Sheets: Trades (all entries), Summary (statistics with formulas)")
    
    doc.add_heading("Excel Columns", level=2)
    
    excel_columns = [
        ("Trade ID", "Unique identifier (PAPER-YYYYMMDD-XXXXXX)"),
        ("Date/Time", "Entry timestamp"),
        ("Symbol", "Trading symbol"),
        ("Action", "BUY or SELL"),
        ("Quantity", "Number of shares"),
        ("Entry Price", "Entry price"),
        ("Exit Price", "Exit price (filled on close)"),
        ("Stop Loss", "Stop loss level"),
        ("Take Profit", "Take profit target"),
        ("Strategy", "AI_COMPOSITE"),
        ("Entry Reason", "Detailed reasoning (e.g., 'Strong technicals (72%); Bull flag (68%)')"),
        ("Signal Confidence", "AI confidence (%)"),
        ("Technical Score", "Technical analysis score (%)"),
        ("Pattern Score", "Pattern detection score (%)"),
        ("Sentiment Score", "News sentiment score (%)"),
        ("P&L ($)", "Dollar profit/loss"),
        ("P&L (%)", "Percentage profit/loss"),
        ("Trade Duration", "Time in position"),
        ("Status", "OPEN or CLOSED"),
        ("Notes", "Additional notes"),
    ]
    
    table = doc.add_table(rows=len(excel_columns) + 1, cols=2)
    table.style = 'Table Grid'
    
    table.rows[0].cells[0].text = "Column"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    for i, (col, desc) in enumerate(excel_columns, 1):
        table.rows[i].cells[0].text = col
        table.rows[i].cells[1].text = desc
    
    doc.add_page_break()
    
    # =====================================================
    # 11. RISK MANAGEMENT
    # =====================================================
    doc.add_heading("11. Risk Management", level=1)
    
    doc.add_heading("Position Sizing Algorithm", level=2)
    
    doc.add_paragraph("The system uses risk-based position sizing:")
    
    sizing_formula = """1. Risk Amount = Portfolio Value × Risk% 
   Example: $100,000 × 2% = $2,000

2. Stop Loss = Entry Price - (ATR × 2)
   Example: $178.50 - ($1.88 × 2) = $174.74

3. Shares = Risk Amount ÷ (Entry - Stop Loss)
   Example: $2,000 ÷ $3.76 = 531 shares

4. Position Value Check = Shares × Entry Price
   Example: 531 × $178.50 = $94,783
   
5. If > MAX_POSITION_SIZE (10%), reduce:
   Shares = ($100,000 × 10%) ÷ $178.50 = 56 shares"""
    
    p = doc.add_paragraph()
    p.add_run(sizing_formula).font.name = 'Courier New'
    
    doc.add_heading("Default Risk Parameters", level=2)
    
    risk_params = [
        ("Parameter", "Default", "Description"),
        ("risk_per_trade", "2%", "Maximum loss per trade"),
        ("DEFAULT_STOP_LOSS", "2%", "Stop loss distance from entry"),
        ("DEFAULT_TAKE_PROFIT", "5%", "Take profit target"),
        ("MAX_POSITION_SIZE", "10%", "Maximum portfolio % per position"),
        ("AUTO_MAX_POSITIONS", "5", "Maximum concurrent positions"),
    ]
    
    table = doc.add_table(rows=len(risk_params), cols=3)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(risk_params):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # =====================================================
    # 12. TROUBLESHOOTING
    # =====================================================
    doc.add_heading("12. Troubleshooting", level=1)
    
    doc.add_heading("Common Issues and Solutions", level=2)
    
    issues = [
        ("'Import could not be resolved' in IDE",
         "This is normal before dependencies are installed.",
         "pip install -r requirements.txt"),
        
        ("Trading not starting automatically",
         "Check AUTO_START_TRADING setting in .env",
         "Set AUTO_START_TRADING=true"),
        
        ("No signals being generated",
         "AI_CONFIDENCE_THRESHOLD may be too high",
         "Try lowering to 0.60"),
        
        ("Excel log file not created",
         "openpyxl library not installed",
         "pip install openpyxl"),
        
        ("WebSocket connection failing",
         "Backend not running or CORS issue",
         "Check CORS_ORIGINS includes frontend URL"),
    ]
    
    for issue, cause, solution in issues:
        doc.add_heading(issue, level=3)
        p = doc.add_paragraph()
        p.add_run("Cause: ").bold = True
        p.add_run(cause)
        p = doc.add_paragraph()
        p.add_run("Solution: ").bold = True
        p.add_run(solution).font.name = 'Courier New'
        doc.add_paragraph()
    
    doc.add_heading("Quick Reference Commands", level=2)
    
    commands = """# Start backend (development)
cd backend
uvicorn app.main:app --reload --host 0.0.0.0 --port 8000

# Start frontend (development)
cd frontend
npm start

# Docker start/stop
docker-compose up -d
docker-compose down

# Check trading status
curl http://localhost:8000/api/auto/status

# Start/Stop trading
curl -X POST http://localhost:8000/api/auto/start
curl -X POST http://localhost:8000/api/auto/stop

# Download trade log
curl http://localhost:8000/api/log/download -o trade_log.xlsx"""
    
    p = doc.add_paragraph()
    p.add_run(commands).font.name = 'Courier New'
    
    # =====================================================
    # FOOTER
    # =====================================================
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Documentation last updated: March 2026")
    footer_run.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    doc.save('AI_Trading_Platform_Documentation.docx')
    print("✅ Word document created: AI_Trading_Platform_Documentation.docx")


if __name__ == "__main__":
    create_trading_platform_doc()
